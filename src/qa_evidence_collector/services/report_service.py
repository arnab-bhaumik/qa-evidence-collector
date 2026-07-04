from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from qa_evidence_collector.core.session_manager import SessionManager


class ReportService:
    def generate(self, session: SessionManager, output_dir: str | Path) -> str:
        doc = Document()
        self._set_page_margins(doc)
        self._build_title_page(doc, session)
        self._build_steps(doc, session)

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        # Use Test Case ID only (short, safe) + sequential number
        base = session.test_case_id if session.test_case_id else session.session_name
        safe_base = "".join(
            c if c.isalnum() or c in "_-" else "_"
            for c in base
        ).strip("_")

        # Find next sequential number in the folder
        existing = list(output_dir.glob(f"{safe_base}_*.docx"))
        next_num = len(existing) + 1
        filename = f"{safe_base}_{next_num:02d}.docx"
        filepath = output_dir / filename

        doc.save(str(filepath))
        return str(filepath)

    # ------------------------------------------------------------------
    # Title page
    # ------------------------------------------------------------------

    def _build_title_page(self, doc: Document, session: SessionManager) -> None:
        doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("QA Evidence Report")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        doc.add_paragraph()

        rows = [
            ("Test Case ID", session.test_case_id or "—"),
            ("Test Objective", session.test_objective or "—"),
            ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("Total Steps", str(len(session.steps))),
        ]

        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.autofit = True

        for i, (label, value) in enumerate(rows):
            label_cell = table.cell(i, 0)
            value_cell = table.cell(i, 1)

            label_run = label_cell.paragraphs[0].add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(11)

            value_run = value_cell.paragraphs[0].add_run(value)
            value_run.font.size = Pt(11)

        doc.add_paragraph()
        doc.add_paragraph().add_run("_" * 60).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        doc.add_page_break()

    # ------------------------------------------------------------------
    # Steps
    # ------------------------------------------------------------------

    def _build_steps(self, doc: Document, session: SessionManager) -> None:
        for step in session.steps:
            # Step heading
            heading = doc.add_paragraph()
            run = heading.add_run(f"Step {step.step_number}")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

            # Note
            note_text = step.note if step.note else "(No note provided)"
            note_para = doc.add_paragraph()
            note_run = note_para.add_run(note_text)
            note_run.font.size = Pt(11)

            # Screenshot
            img_path = Path(step.screenshot_path)
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(6.0))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph("[Screenshot could not be embedded]")
            else:
                doc.add_paragraph(f"[Screenshot not found: {img_path}]")

            doc.add_paragraph()
            doc.add_paragraph().add_run("─" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            doc.add_paragraph()

    # ------------------------------------------------------------------
    # Page margins
    # ------------------------------------------------------------------

    def _set_page_margins(self, doc: Document) -> None:
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
